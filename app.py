from flask import Flask, render_template, request, send_file, flash, send_from_directory, jsonify
from openpyxl import load_workbook
from docx import Document
import os
from io import BytesIO
import logging
from logging.handlers import RotatingFileHandler
import zipfile
import re
from datetime import datetime
import json

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'your-secret-key-here')

# Настройка логирования
if not os.path.exists('logs'):
    os.makedirs('logs')
file_handler = RotatingFileHandler('logs/kp_generator.log', maxBytes=10240, backupCount=10)
file_handler.setFormatter(logging.Formatter(
    '%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]'))
file_handler.setLevel(logging.INFO)
app.logger.addHandler(file_handler)
app.logger.setLevel(logging.INFO)
app.logger.info('KP Generator startup')

def validate_form_data(form_data):
    """Проверяет корректность данных формы"""
    errors = []
    
    # Проверка обязательных полей
    required_fields = ['company', 'product', 'quantity', 'cost_price', 'weight', 'logistics']
    for field in required_fields:
        if not form_data.get(field) or not form_data[field].strip():
            errors.append(f'Поле "{field}" является обязательным.')
    
    # Проверка числовых значений
    numeric_fields = ['quantity', 'cost_price', 'weight', 'logistics', 'duty_percent', 'deal_length_days']
    for field in numeric_fields:
        if form_data.get(field) and form_data[field].strip():
            try:
                value = float(form_data[field])
                if value < 0:
                    errors.append(f'Поле "{field}" должно быть неотрицательным числом.')
                if field == 'duty_percent' and value > 100:
                    errors.append(f'Поле "{field}" не может превышать 100%.')
                if field == 'quantity' and value == 0:
                    errors.append(f'Поле "{field}" не может быть нулевым.')
                if field == 'deal_length_days' and value < 30:
                    errors.append(f'Поле "{field}" не может быть меньше 30 дней.')
            except ValueError:
                errors.append(f'Поле "{field}" должно быть числом.')
    
    return errors

def get_safe_filename(company_name):
    """Создает безопасное имя файла из названия компании"""
    safe_name = re.sub(r'[^\w\s-]', '', company_name).strip()
    safe_name = re.sub(r'[-\s]+', '_', safe_name)
    return safe_name[:50]

def calculate_selling_price(quantity, purchase_cost, logistics_rub, duty_percent, weight, deal_length_days=170, margin_percent=30):
    """Выполняет расчет продажной цены с учетом всех параметров бюджета"""
    # Константы из бюджета
    CONVERSION_RATE = 12  # Курс юаня к рублю
    LOGISTICS_CNR_RATIO = 0.3  # Доля логистики КНР
    LOGISTICS_RF_RATIO = 0.7  # Доля логистики РФ
    CONVERSION_FEE_RATE = 0.032  # Комиссия за конвертацию 3.2%
    CREDIT_RATE = 0.16  # Ставка кредита 16%
    
    # Расчет общего веса
    total_weight = weight * quantity
    
    # Перевод логистики в юани и распределение по весу
    logistics_total_yuan = logistics_rub / CONVERSION_RATE
    
    # Расчет логистики на единицу товара (пропорционально весу)
    logistics_cnr_per_unit = (logistics_total_yuan * LOGISTICS_CNR_RATIO * weight) / total_weight
    logistics_rf_per_unit = (logistics_total_yuan * LOGISTICS_RF_RATIO * weight) / total_weight
    
    # Расчет пошлины на единицу товара
    duty_per_unit = (purchase_cost + logistics_cnr_per_unit) * (duty_percent / 100)
    
    # Расчет стоимости конвертации
    conversion_fee = purchase_cost * quantity * CONVERSION_FEE_RATE
    conversion_fee_per_unit = conversion_fee / quantity
    
    # Расчет кредитных затрат
    credit_cost = purchase_cost * quantity * CREDIT_RATE / 365 * deal_length_days
    credit_cost_per_unit = credit_cost / quantity
    
    # Общие затраты на единицу товара
    total_cost_per_unit = (
        purchase_cost +
        logistics_cnr_per_unit +
        logistics_rf_per_unit +
        duty_per_unit +
        conversion_fee_per_unit +
        credit_cost_per_unit
    )
    
    # Расчет цены для маржи margin_percent%
    selling_price_per_unit = total_cost_per_unit / (1 - margin_percent / 100)
    
    return selling_price_per_unit

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/favicon.ico')
def favicon():
    return send_from_directory(os.path.join(app.root_path, 'static'),
                               'favicon.ico', mimetype='image/vnd.microsoft.icon')

@app.route('/generate', methods=['POST'])
def generate():
    form_data = request.form.to_dict()
    errors = validate_form_data(form_data)
    
    if errors:
        for error in errors:
            flash(error, 'danger')
        return render_template('index.html', form_data=form_data)
    
    try:
        # Извлечение и преобразование данных
        company = form_data['company'].strip()
        product = form_data['product'].strip()
        quantity = int(form_data['quantity'])
        cost_price = float(form_data['cost_price'])
        weight = float(form_data['weight'])
        logistics = float(form_data['logistics'])
        
        # Новые поля
        tender_number = form_data.get('tender_number', '').strip()
        drawing_number = form_data.get('drawing_number', '').strip()
        material = form_data.get('material', '').strip()
        delivery_address = form_data.get('delivery_address', '').strip()
        duty_percent = float(form_data.get('duty_percent', 0))
        deal_length_days = float(form_data.get('deal_length_days', 170))
        
        # Расчет сроков поставки и оплаты
        supply_days = deal_length_days - 30
        payment_days = 30
        
        if supply_days < 0:
            flash('Общая длина сделки не может быть меньше 30 дней.', 'danger')
            return render_template('index.html', form_data=form_data)
        
        # Формируем текст для ячейки C10
        product_with_drawing = product
        if drawing_number:
            product_with_drawing += f" ч.{drawing_number}"
        
        # Выполнение расчетов с учетом всех параметров бюджета
        final_price = calculate_selling_price(
            quantity=quantity,
            purchase_cost=cost_price,
            logistics_rub=logistics,
            duty_percent=duty_percent,
            weight=weight,
            deal_length_days=deal_length_days,
            margin_percent=30  # Целевая маржа 30%
        )
        
        # Проверка существования шаблонов
        excel_template_path = os.path.join('templates_docs', 'template.xlsx')
        word_template_path = os.path.join('templates_docs', 'template.docx')
        
        if not os.path.exists(excel_template_path):
            flash('Шаблон Excel не найден. Обратитесь к администратору.', 'danger')
            app.logger.error(f'Excel template not found: {excel_template_path}')
            return render_template('index.html', form_data=form_data)
            
        if not os.path.exists(word_template_path):
            flash('Шаблон Word не найден. Обратитесь к администратору.', 'danger')
            app.logger.error(f'Word template not found: {word_template_path}')
            return render_template('index.html', form_data=form_data)
        
        # Работа с Excel
        try:
            wb = load_workbook(excel_template_path)
            ws = wb.active
            
            # Заполняем данные в Excel
            current_date = datetime.now().strftime('%d.%m.%Yг.')
            
            # Основные данные
            ws['D4'] = company
            ws['C10'] = product_with_drawing  # Наименование товара с номером чертежа
            ws['G10'] = quantity
            ws['M10'] = cost_price
            ws['P10'] = weight
            ws['U14'] = logistics
            ws['X10'] = duty_percent / 100  # Процент пошлины (доля)
            ws['I15'] = supply_days  # Срок поставки
            ws['I16'] = payment_days  # Срок оплаты
            
            # Дополнительные поля
            ws['D2'] = current_date  # Дата формирования
            ws['D5'] = tender_number  # Номер тендера
            ws['D10'] = material  # Материал
            ws['E10'] = drawing_number  # Номер чертежа
            ws['P4'] = delivery_address  # Адрес доставки
            
            # Расчетные поля
            ws['H10'] = final_price  # Финальная цена
            
            excel_file = BytesIO()
            wb.save(excel_file)
            excel_file.seek(0)
        except Exception as e:
            flash('Ошибка при обработке Excel-шаблона.', 'danger')
            app.logger.error(f'Excel processing error: {str(e)}')
            return render_template('index.html', form_data=form_data)
        
        # Работа с Word
        try:
            doc = Document(word_template_path)
            
            # Форматируем текущую дату
            current_date = datetime.now().strftime('%d.%m.%Yг.')
            
            word_data = {
                '{{ company }}': company,
                '{{ product }}': product,
                '{{ quantity }}': str(quantity),
                '{{ cost_price }}': f"{cost_price:.2f}",
                '{{ weight }}': f"{weight:.2f}",
                '{{ logistics }}': f"{logistics:.2f}",
                '{{ final_price }}': f"{final_price:.2f}",
                '{{ tender_number }}': tender_number,
                '{{ drawing_number }}': drawing_number,
                '{{ material }}': material,
                '{{ delivery_address }}': delivery_address,
                '{{ date }}': current_date,
                '{{ duty_percent }}': f"{duty_percent:.1f}",
                '{{ deal_length_days }}': str(deal_length_days),
                '{{ supply_days }}': str(supply_days),
                '{{ payment_days }}': str(payment_days),
            }
            
            # Замена в параграфах
            for paragraph in doc.paragraphs:
                for key, value in word_data.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)
            
            # Замена в таблицах
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in word_data.items():
                            if key in cell.text:
                                cell.text = cell.text.replace(key, value)
            
            word_file = BytesIO()
            doc.save(word_file)
            word_file.seek(0)
        except Exception as e:
            flash('Ошибка при обработке Word-шаблона.', 'danger')
            app.logger.error(f'Word processing error: {str(e)}')
            return render_template('index.html', form_data=form_data)
        
        # Создание ZIP-архива
        file_prefix = f"КП_{get_safe_filename(company)}_{datetime.now().strftime('%Y%m%d_%H%M')}"
        
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Добавляем Excel-файл
            zip_file.writestr(f"{file_prefix}.xlsx", excel_file.getvalue())
            # Добавляем Word-файл
            zip_file.writestr(f"{file_prefix}.docx", word_file.getvalue())
        
        zip_buffer.seek(0)
        
        return send_file(
            zip_buffer,
            as_attachment=True,
            download_name=f"{file_prefix}.zip",
            mimetype='application/zip'
        )
    
    except Exception as e:
        flash('Произошла непредвиденная ошибка. Попробуйте еще раз.', 'danger')
        app.logger.error(f'Unexpected error: {str(e)}')
        return render_template('index.html', form_data=form_data)

@app.errorhandler(404)
def not_found_error(error):
    return render_template('404.html'), 404

@app.errorhandler(500)
def internal_error(error):
    return render_template('500.html'), 500

if __name__ == '__main__':
    # Создаем необходимые папки при запуске
    for folder in ['logs', 'templates_docs']:
        if not os.path.exists(folder):
            os.makedirs(folder)
    
    app.run(debug=True, host='0.0.0.0', port=5000)