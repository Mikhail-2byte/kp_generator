"""
Скрипт для анализа структуры Word шаблона.
Помогает понять, где находятся таблицы, строки и плейсхолдеры.
"""
from docx import Document
from docx.shared import Inches
import sys
import os

# Добавляем корневую директорию в путь
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

def analyze_word_template(template_path):
    """Анализирует структуру Word шаблона."""
    print(f"Анализ шаблона: {template_path}\n")
    print("=" * 80)
    
    if not os.path.exists(template_path):
        print(f"ОШИБКА: Файл {template_path} не найден!")
        return
    
    doc = Document(template_path)
    
    # Анализ параграфов
    print("\n=== ПАРАГРАФЫ ===")
    for i, paragraph in enumerate(doc.paragraphs, 1):
        text = paragraph.text.strip()
        if text:
            # Ищем плейсхолдеры
            placeholders = []
            if '{{' in text and '}}' in text:
                import re
                placeholders = re.findall(r'\{\{[^}]+\}\}', text)
            
            if placeholders or len(text) < 100:
                print(f"\nПараграф {i}:")
                print(f"  Текст: {text[:200]}")
                if placeholders:
                    print(f"  Плейсхолдеры: {placeholders}")
    
    # Анализ таблиц
    print("\n\n=== ТАБЛИЦЫ ===")
    print(f"Всего таблиц: {len(doc.tables)}\n")
    
    for table_idx, table in enumerate(doc.tables, 1):
        print(f"\n--- Таблица {table_idx} ---")
        print(f"Строк: {len(table.rows)}, Столбцов: {len(table.columns)}")
        
        # Анализируем каждую строку
        for row_idx, row in enumerate(table.rows, 1):
            row_data = []
            has_placeholder = False
            
            for cell_idx, cell in enumerate(row.cells, 1):
                cell_text = cell.text.strip()
                if '{{' in cell_text and '}}' in cell_text:
                    has_placeholder = True
                    import re
                    placeholders = re.findall(r'\{\{[^}]+\}\}', cell_text)
                    row_data.append(f"[{cell_idx}] {cell_text[:50]}... (плейсхолдеры: {placeholders})")
                elif cell_text:
                    row_data.append(f"[{cell_idx}] {cell_text[:50]}")
                else:
                    row_data.append(f"[{cell_idx}] (пусто)")
            
            if has_placeholder or any(data for data in row_data if data and not data.endswith("(пусто)")):
                print(f"\n  Строка {row_idx}:")
                for data in row_data:
                    if data and not data.endswith("(пусто)"):
                        print(f"    {data}")
    
    print("\n" + "=" * 80)
    print("\nАнализ завершен!")

if __name__ == "__main__":
    template_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'templates_docs', 'template.docx')
    analyze_word_template(template_path)

