"""
Простой скрипт для проверки структуры Word шаблона.
Запустите в окружении с установленным python-docx.
"""
import sys
import os

# Добавляем корневую директорию в путь
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

try:
    from docx import Document
except ImportError:
    print("❌ Ошибка: python-docx не установлен")
    print("Установите: pip install python-docx")
    sys.exit(1)

def check_template():
    """Проверяет структуру шаблона."""
    template_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'templates_docs', 'template.docx')
    
    if not os.path.exists(template_path):
        print(f"❌ Шаблон не найден: {template_path}")
        return
    
    print(f"Проверка шаблона: {template_path}\n")
    print("=" * 70)
    
    doc = Document(template_path)
    
    # Проверяем таблицы
    print(f"\n📊 Найдено таблиц: {len(doc.tables)}")
    
    position_placeholders = ['{{ product }}', '{{ quantity }}', '{{ cost_price }}', '{{ weight }}']
    common_placeholders = ['{{ company }}', '{{ date }}', '{{ logistics }}', '{{ delivery_time }}']
    
    found_position_table = False
    found_template_row = False
    
    for table_idx, table in enumerate(doc.tables, 1):
        print(f"\n--- Таблица {table_idx} ---")
        print(f"   Строк: {len(table.rows)}, Столбцов: {len(table.columns)}")
        
        has_position_placeholders = False
        template_row_idx = None
        
        for row_idx, row in enumerate(table.rows):
            row_text = ' '.join(cell.text for cell in row.cells)
            
            # Проверяем наличие плейсхолдеров позиций
            found_placeholders = [ph for ph in position_placeholders if ph in row_text]
            if found_placeholders:
                has_position_placeholders = True
                if template_row_idx is None:
                    template_row_idx = row_idx
                print(f"   ✅ Строка {row_idx + 1}: найдены плейсхолдеры позиций: {found_placeholders}")
        
        if has_position_placeholders:
            found_position_table = True
            found_template_row = True
            print(f"   ✅ Это таблица с позициями! Строка-шаблон: строка {template_row_idx + 1}")
    
    # Проверяем общие плейсхолдеры в параграфах
    print(f"\n📝 Проверка параграфов:")
    found_common = []
    for para_idx, para in enumerate(doc.paragraphs[:20], 1):  # Первые 20 параграфов
        para_text = para.text
        for ph in common_placeholders:
            if ph in para_text and ph not in found_common:
                found_common.append(ph)
                print(f"   ✅ Найден плейсхолдер: {ph}")
    
    # Итоговая проверка
    print("\n" + "=" * 70)
    print("ИТОГОВАЯ ПРОВЕРКА:")
    print("=" * 70)
    
    if found_position_table:
        print("✅ Таблица с позициями найдена")
    else:
        print("❌ Таблица с позициями НЕ найдена")
        print("   Нужно добавить таблицу со строкой, содержащей плейсхолдеры:")
        print("   {{ product }}, {{ quantity }}, {{ cost_price }}, {{ weight }}")
    
    if found_template_row:
        print("✅ Строка-шаблон найдена")
    else:
        print("❌ Строка-шаблон НЕ найдена")
    
    if found_common:
        print(f"✅ Найдено общих плейсхолдеров: {len(found_common)}")
    else:
        print("⚠️  Общие плейсхолдеры не найдены в первых 20 параграфах")
    
    print("\n" + "=" * 70)
    print("РЕКОМЕНДАЦИИ:")
    print("=" * 70)
    
    if not found_position_table:
        print("""
1. Создайте таблицу в Word шаблоне
2. Добавьте строку-шаблон с плейсхолдерами позиций:
   - {{ product }} - наименование товара
   - {{ quantity }} - количество
   - {{ cost_price }} - сумма закупа
   - {{ weight }} - вес за шт
   - {{ drawing_number }} - номер чертежа (опционально)
   - {{ material }} - материал (опционально)
   - {{ duty_percent }} - пошлина % (опционально)
   
3. Первая строка таблицы может быть заголовком
4. Вторая строка (или первая, если нет заголовка) должна содержать плейсхолдеры
        """)
    else:
        print("✅ Шаблон готов к использованию!")
        print("   Процессор автоматически найдет таблицу и строку-шаблон")

if __name__ == "__main__":
    check_template()

